import io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generate_bid_proposal_docx(match_data: dict, vendor_profile: dict) -> io.BytesIO:
    """
    Generates a pre-filled compliance matrix and bid proposal template
    based on the AI Match data and the Vendor Profile.
    """
    document = Document()

    # --- Title Page ---
    title = document.add_heading('Bid Proposal & Compliance Matrix', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph(f"Tender Name: {match_data.get('tender_title', 'Unknown Tender')}")
    document.add_paragraph(f"Tender ID: {match_data.get('tender_id', 'N/A')}")
    
    vendor_name = vendor_profile.get("profile_data", {}).get("identity", {}).get("company_legal_name", "Vendor")
    document.add_paragraph(f"Prepared By: {vendor_name}")
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    document.add_page_break()

    # --- Section 1: Executive Summary ---
    document.add_heading('1. Executive Summary', level=1)
    exec_summary = document.add_paragraph(
        f"Based on the AI matching analysis, {vendor_name} is considered a "
    )
    run_rec = exec_summary.add_run(match_data.get('recommendation', 'Candidate'))
    run_rec.bold = True
    
    if match_data.get("recommendation", "") == "Strongly Recommended":
        run_rec.font.color.rgb = RGBColor(0, 128, 0) # Green
    
    exec_summary.add_run(f" for this tender with an overall algorithmic match score of {match_data.get('final_score', 0)}/100.")
    
    document.add_heading('Score Breakdown:', level=2)
    breakdown = match_data.get('score_breakdown', {})
    for key, val in breakdown.items():
        document.add_paragraph(f"• {key.replace('_', ' ').title()}: {val}%", style='List Bullet')
        
    # --- Section 2: AI Risk Analysis (Devil's Advocate) ---
    if "devil_advocate" in match_data:
        document.add_heading('2. Risk Analysis & Mitigation Strategy', level=1)
        da = match_data["devil_advocate"]
        document.add_paragraph(da.get("summary", ""))
        
        document.add_heading('Identified Risks to Mitigate in Proposal:', level=2)
        for risk in da.get("critical_risks", []):
            document.add_paragraph(risk, style='List Bullet')
            document.add_paragraph("Proposed Mitigation: [Insert Mitigation Strategy Here]")

    # --- Section 3: Compliance Matrix ---
    document.add_heading('3. Compliance Matrix', level=1)
    
    # Create Table
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Requirement Category'
    hdr_cells[1].text = 'Tender Requirement (Auto-extracted)'
    hdr_cells[2].text = 'Vendor Compliance / Response'
    
    # Populate Table rows based on extracted data
    # (Mocking some standard rows, in production this would map to exact extracted JSON)
    
    rows_data = [
        ("Financial Turnover", "Meet minimum turnover requirements", "COMPLIANT" if match_data.get("is_eligible") else "NON-COMPLIANT"),
        ("Domain Expertise", f"Sector: {match_data.get('sector', 'N/A')}", "COMPLIANT"),
        ("Geographic Presence", "Operational in target region", "COMPLIANT"),
        ("Certifications", "Required certifications met", "COMPLIANT"),
    ]
    
    for category, req, comp in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = req
        row_cells[2].text = comp

    document.add_page_break()

    # --- Section 4: Draft Proposal Template ---
    document.add_heading('4. Technical Proposal Draft', level=1)
    document.add_paragraph("[This is an AI-generated draft structure. Please expand.]")
    document.add_heading('4.1 Understanding of Scope', level=2)
    document.add_paragraph(f"We understand the core requirement is related to {match_data.get('sector', 'the specified sector')}...")
    document.add_heading('4.2 Methodology', level=2)
    document.add_paragraph("[Insert Methodology Here]")
    document.add_heading('4.3 Past Experience', level=2)
    document.add_paragraph(f"We have extensive experience in similar projects as detailed in our profile...")

    # Save to memory
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
